from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def add_heading_with_color(doc, text, level, color=RGBColor(0, 51, 102)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def main():
    project_root = Path(r'd:\gemma4\segmentation_lattuce-desease')
    assets_dir = project_root / 'Report' / 'stage2_assets'
    artifact_dir = Path(r'C:\Users\amans\.gemini\antigravity\brain\1b812dbc-7162-40d7-84ab-ca5bf058a0b9')
    output_path = project_root / 'Report' / 'Stage2_Healthy_Learning_Technical_Documentation.docx'

    doc = Document()

    # --- Title Page ---
    title = doc.add_heading('Stage 2: Healthy-Only Representation Learning', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Technical Deep Dive into Self-Supervised Feature Extraction and Statistical Modeling')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- 1. Executive Summary ---
    add_heading_with_color(doc, '1. Executive Summary', level=1)
    doc.add_paragraph(
        "Stage 2 of the Lettuce Disease Segmentation pipeline focuses on learning a robust 'normality' model using self-supervised representations. "
        "By analyzing only healthy lettuce leaves, the system builds a statistical 'Healthy Bank' that defines the expected distribution of visual features. "
        "This approach allows for anomaly-based localization in Stage 3, where any significant deviation from this learned normality is flagged as a potential disease region."
    )

    # --- 2. Architecture: DINOv2 & Vision Transformers ---
    add_heading_with_color(doc, '2. Architecture: DINOv2 & ViT-B/14', level=1)
    doc.add_paragraph(
        "The backbone of this stage is DINOv2 (Self-DIstillation with NO labels), a foundation model trained by Meta AI. "
        "Unlike supervised models, DINOv2 learns features that are highly semantic and geometrically robust without needing manual annotations."
    )
    
    # Insert DINOv2 Architecture Diagram
    dinov2_diagram = list(artifact_dir.glob('dinov2_architecture_diagram_*.png'))
    if dinov2_diagram:
        doc.add_picture(str(dinov2_diagram[0]), width=Inches(5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 1: DINOv2 Vision Transformer Architecture showing patch-level feature extraction.').alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('Why DINOv2?', level=2)
    doc.add_paragraph(
        "1. All-Purpose Features: Performs exceptionally well on segmentation tasks out-of-the-box.\n"
        "2. Patch Awareness: Provides dense, local features rather than just a global image embedding.\n"
        "3. Invariance: Robust to changes in lighting and leaf orientation."
    )

    # --- 3. Data Pipeline and Tensor Shapes ---
    add_heading_with_color(doc, '3. Data Pipeline and Tensor Shapes', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Stage'
    hdr_cells[1].text = 'Operation'
    hdr_cells[2].text = 'Tensor Shape (B, C, H, W)'
    
    steps = [
        ('Input', 'MultiChannel Dataset (RGB Slice)', '(B, 3, 256, 256)'),
        ('Preprocessing', 'Resize to patch multiple (14px)', '(B, 3, 252, 252)'),
        ('Extraction', 'DINOv2 ViT-B/14 forward pass', '(B, 768, 18, 18)'),
        ('Reduction', 'Random Dimensionality Selection', '(B, 128, 18, 18)'),
        ('Modeling', 'Gaussian Parameter Estimation', '(18*18, 128, 128)')
    ]
    
    for stage, op, shape in steps:
        row_cells = table.add_row().cells
        row_cells[0].text = stage
        row_cells[1].text = op
        row_cells[2].text = shape

    # --- 4. Mathematical Methodology: PaDiM ---
    add_heading_with_color(doc, '4. Mathematical Methodology: PaDiM', level=1)
    doc.add_paragraph(
        "We utilize Patch Distribution Modeling (PaDiM) to capture the statistical structure of healthy leaves. "
        "For every spatial patch position 'p' in the 18x18 grid, we model the distribution of features as a Multivariate Gaussian."
    )
    
    # Insert PaDiM Workflow Diagram
    padim_diagram = list(artifact_dir.glob('padim_workflow_diagram_*.png'))
    if padim_diagram:
        doc.add_picture(str(padim_diagram[0]), width=Inches(5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 2: PaDiM Workflow - Multiple images contribute to a local Gaussian distribution per patch.').alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('The Gaussian Distribution', level=2)
    doc.add_paragraph(
        "For each patch position, we calculate the Sample Mean (μ) and Covariance Matrix (Σ):"
    )
    doc.add_paragraph("μ_p = (1/N) * Σ(x_p,k)")
    doc.add_paragraph("Σ_p = (1/(N-1)) * Σ(x_p,k - μ_p)(x_p,k - μ_p)^T + εI")
    
    doc.add_paragraph(
        "Where 'N' is the number of healthy images and 'εI' is a regularization term used to ensure the covariance matrix is invertible."
    )

    # --- 5. Visual Evidence: Healthy Features ---
    add_heading_with_color(doc, '5. Visual Evidence: Healthy Features', level=1)
    
    # Activation Map
    act_map = assets_dir / '01_healthy_activation.png'
    if act_map.exists():
        doc.add_picture(str(act_map), width=Inches(5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 3: DINOv2 Activation Map showing high sensitivity to leaf structure and texture.').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Channels
    chan_map = assets_dir / '02_feature_channels.png'
    if chan_map.exists():
        doc.add_picture(str(chan_map), width=Inches(5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 4: Individual feature channels highlighting different spatial and spectral properties of the leaf.').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 6. Conclusion ---
    add_heading_with_color(doc, '6. Conclusion', level=1)
    doc.add_paragraph(
        "By the end of Stage 2, the system has constructed a high-dimensional statistical profile of healthy lettuce. "
        "The resulting 'Healthy Bank' (stored as .pkl and .npy files) acts as a reference library for all future anomaly detection tasks, "
        "providing a mathematically rigorous definition of 'normal leaf' against which diseased samples will be compared."
    )

    doc.save(str(output_path))
    print(f"Report saved to {output_path}")

if __name__ == '__main__':
    main()
