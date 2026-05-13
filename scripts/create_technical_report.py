from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def add_step(doc, title, image_path, image_desc, significance_ip, significance_disease):
    doc.add_heading(title, level=2)
    
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(3))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.add_run('Description: ').bold = True
    p.add_run(image_desc)
    
    p = doc.add_paragraph()
    p.add_run('Significance in Image Processing: ').bold = True
    p.add_run(significance_ip)
    
    p = doc.add_paragraph()
    p.add_run('Significance in Disease Region Capturing: ').bold = True
    p.add_run(significance_disease)

def main():
    project_root = Path(r'd:\gemma4\segmentation_lattuce-desease')
    assets_dir = project_root / 'Report' / 'preprocessing_steps'
    output_path = project_root / 'Report' / 'Preprocessing_Methodology.docx'

    doc = Document()

    # Title
    title = doc.add_heading('Technical Documentation: Preprocessing Methodology', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('This document outlines the multi-stage preprocessing pipeline developed for the self-supervised lettuce disease segmentation system. Each step is designed to transform the raw RGB input into specialized representations that enhance the model\'s ability to localize and segment diseased regions accurately.')

    # 1. Original RGB
    add_step(doc, '1. Original RGB Input', 
             assets_dir / '01_original_rgb.png',
             'The raw input image resized to 256x256 pixels.',
             'Serves as the baseline visual representation, containing all spectral information available from the sensor.',
             'Provides the primary context for the lesion, though raw images often suffer from lighting variations and low contrast in diseased spots.')

    # 2. CLAHE
    add_step(doc, '2. Contrast Limited Adaptive Histogram Equalization (CLAHE)', 
             assets_dir / '02_clahe.png',
             'The image after applying CLAHE in the LAB color space (L-channel).',
             'Enhances local contrast by redistributing pixel intensities within small tiles, preventing over-amplification of noise while making details pop.',
             'Critical for identifying "halo" effects or subtle color gradients at the edges of bacterial spots that might be washed out in standard RGB.')

    # 3. Excess Green Index (ExG)
    add_step(doc, '3. Excess Green Index (ExG)', 
             assets_dir / '03_exg.png',
             'A vegetation index map calculated as 2G - R - B.',
             'Highlights chlorophyll-rich regions by emphasizing the green channel relative to red and blue.',
             'Specifically useful for background subtraction and identifying necrotic (dead) tissue, which typically shows a sharp drop in ExG values compared to healthy green leaves.')

    # 4. Edge Map (Sobel)
    add_step(doc, '4. Edge Detection (Sobel Gradient)', 
             assets_dir / '04_edge_map.png',
             'A magnitude map of horizontal and vertical gradients.',
             'Captures high-frequency spatial information, highlighting boundaries and texture transitions.',
             'Diseased regions often have distinct, irregular boundaries. The edge map helps the model focus on these structural discontinuities during segmentation.')

    # 5. Felzenszwalb Segmentation (Superpixels)
    add_step(doc, '5. Felzenszwalb Superpixel Segmentation', 
             assets_dir / '06_felz_colored.png',
             'A graph-based segmentation that groups pixels into homogeneous regions.',
             'Reduces the complexity of the image from thousands of pixels to a few hundred "superpixels" that preserve boundary information.',
             'Ensures that entire lesion areas are treated as cohesive units rather than isolated pixels, facilitating more consistent region-based localization.')

    # 6. Watershed Segmentation
    add_step(doc, '6. Watershed Segmentation', 
             assets_dir / '08_watershed.png',
             'A region-growing segmentation based on the distance transform and image gradients.',
             'Treats the image as a topographic surface where boundaries are "ridges" between "catchment basins".',
             'Excellent for separating overlapping leaves and localizing individual infection centers (foci) within a larger leaf area.')

    # Conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph('By stacking these 14 representations (RGB, CLAHE, Felz, Watershed, etc.), the system provides the neural network with a "multi-modal" view of the same leaf. This methodology significantly reduces the reliance on manual annotations by providing unsupervised structural and spectral priors that guide the anomaly localization and pseudo-mask generation stages.')

    doc.save(str(output_path))
    print(f"Report saved to {output_path}")

if __name__ == '__main__':
    main()
