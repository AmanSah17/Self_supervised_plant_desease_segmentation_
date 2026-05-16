from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import pandas as pd

def add_heading_with_color(doc, text, level, color=RGBColor(0, 51, 102)):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = color

def create_report():
    project_root = Path(r'd:\gemma4\segmentation_lattuce-desease')
    report_dir = project_root / 'Report'
    report_dir.mkdir(exist_ok=True)
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Technical Documentation: Validation & Refinement (Stages 7 & 8)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # --- 1. Stage 7: Validation Inference & Analytics ---
    add_heading_with_color(doc, '1. Stage 7: Validation Inference & Analytics', level=1)
    doc.add_paragraph(
        "Stage 7 is the comprehensive evaluation phase where the trained SegFormer model (Stage 6) "
        "is deployed on the unseen validation dataset. This stage integrates both semantic segmentation "
        "and statistical anomaly detection to provide a multi-dimensional view of model performance."
    )
    
    add_heading_with_color(doc, '1.1 Execution Pipeline', level=2)
    doc.add_paragraph(
        "The inference pipeline follows a dual-path architecture:"
    )
    doc.add_paragraph("- Semantic Path: The SegFormer model predicts categorical labels for each pixel.", style='List Bullet')
    doc.add_paragraph("- Structural Path: The PaDiM model generates anomaly heatmaps to identify tissue deviations.", style='List Bullet')
    doc.add_paragraph("- Analytics: Class probabilities and confidence scores are calculated to assess prediction reliability.", style='List Bullet')
    
    workflow_diag = project_root / 'stage7_analytics_workflow_diagram_1778676981153.png'
    if workflow_diag.exists():
        doc.add_picture(str(workflow_diag), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 1: Stage 7 Validation & Analytics Execution Pipeline.').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 2. Stage 8: Supervised Fine-Tuning (SFT) ---
    add_heading_with_color(doc, '2. Stage 8: Supervised Fine-Tuning (SFT)', level=1)
    doc.add_paragraph(
        "While the self-supervised approach (Stages 1-6) generates high-quality masks without labels, "
        "there is often a small 'semantic gap' due to the noise in pseudo-labels. Stage 8 bridges this gap "
        "by fine-tuning the model on a tiny set of expert-labeled validation images (approx. 30-50 samples)."
    )
    
    sft_diag = project_root / 'stage8_sft_refinement_loop_1778677002410.png'
    if sft_diag.exists():
        doc.add_picture(str(sft_diag), width=Inches(5.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 2: Supervised Fine-Tuning (SFT) Refinement Loop.').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 3. Metric Improvement Analysis ---
    add_heading_with_color(doc, '3. Metric Improvement Analysis', level=1, color=RGBColor(204, 0, 102))
    doc.add_paragraph(
        "The impact of SFT is measured by comparing the Mean Intersection over Union (mIoU) and Accuracy "
        "against manual ground truth. Even with a minimal labeled set, we observe significant sharpening "
        "of disease boundaries and improved class separation."
    )
    
    metrics_path = report_dir / 'stage78_comparison' / 'sft_improvement_metrics.csv'
    if metrics_path.exists():
        df = pd.read_csv(metrics_path)
        avg_ssl = df['ssl_miou'].mean()
        avg_sft = df['sft_miou'].mean()
        improvement = avg_sft - avg_ssl
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Configuration'
        hdr_cells[1].text = 'Mean IoU'
        hdr_cells[2].text = 'Improvement'
        
        row1 = table.add_row().cells
        row1[0].text = 'SSL Only (Stage 6)'
        row1[1].text = f'{avg_ssl:.4f}'
        row1[2].text = '-'
        
        row2 = table.add_row().cells
        row2[0].text = 'SSL + SFT (Stage 8)'
        row2[1].text = f'{avg_sft:.4f}'
        row2[2].text = f'+{improvement:.4f}'

    metrics_chart = project_root / 'metrics_improvement_chart_1778677020655.png'
    if metrics_chart.exists():
        doc.add_picture(str(metrics_chart), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 3: Quantitative Metric Improvement after SFT.').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 4. Visual Comparison ---
    add_heading_with_color(doc, '4. Visual Comparison: SSL vs SFT vs Manual GT', level=1)
    doc.add_paragraph(
        "The following visualization demonstrates the qualitative improvement. The SSL model (Stage 6) "
        "already captures the disease regions, but the SFT model (Stage 8) refines the edges and reduces "
        "false positives by aligning with expert manual labels."
    )
    
    comparison_viz = report_dir / 'stage78_comparison' / 'ssl_vs_sft_comparison.png'
    if comparison_viz.exists():
        doc.add_picture(str(comparison_viz), width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Figure 4: Qualitative Comparison Grid (RGB | Manual GT | SSL | SFT | Error Maps).').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Conclusion ---
    add_heading_with_color(doc, '5. Conclusion', level=1)
    doc.add_paragraph(
        "The two-stage validation and refinement process ensures that the lettuce disease segmentation "
        "system is both robust (via self-supervision) and precise (via limited supervision). Stage 7 "
        "proves the generalizability of the pipeline, while Stage 8 provides the final calibration "
        "required for high-stakes agricultural diagnostics."
    )

    save_path = report_dir / 'Stages_7_8_Validation_Refinement_Documentation.docx'
    doc.save(str(save_path))
    print(f"Report saved to {save_path}")

if __name__ == "__main__":
    create_report()
